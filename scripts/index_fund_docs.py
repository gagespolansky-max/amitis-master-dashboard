#!/usr/bin/env python3
"""Index fund documents into Supabase pgvector for Fund Doc Search.

Run from the master-dashboard project root.
"""

from __future__ import annotations

import argparse
import base64
import contextlib
import csv
import fcntl
import getpass
import hashlib
import json
import os
import re
import sqlite3
import sys
import time
import tempfile
import uuid
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable
from urllib import error, request
from urllib.parse import urlencode

try:
    from dotenv import load_dotenv
except Exception:  # pragma: no cover - optional dependency fallback
    load_dotenv = None

try:
    from tqdm import tqdm
except Exception:  # pragma: no cover - optional dependency fallback
    tqdm = None


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = PROJECT_ROOT / "data"
LOG_DB_PATH = DATA_DIR / "fund_indexing_log.db"
LOCK_DIR = DATA_DIR / "locks"

SUBSCRIPTIONS_ROOT_GRANDLINE = Path(
    "/Users/gage/Amitis Group Holding Dropbox/Amitis Capital - General/AC - Digital/ACDAM Portfolio/Portfolio - Subscriptions/GrandLine"
)
SUBSCRIPTIONS_ROOT_GRANDLINE_FALLBACK = Path(
    "/Users/gage/Amitis Group Holding Dropbox/Amitis Capital - General/AC - Digital/Portfolio - Subscriptions/GrandLine"
)
MANAGER_MATERIALS_ROOT_GRANDLINE = Path(
    "/Users/gage/Amitis Group Holding Dropbox/Amitis Capital - General/AC - Digital/Manager Materials/Grandline"
)
SUBSCRIPTIONS_BASE = Path(
    "/Users/gage/Amitis Group Holding Dropbox/Amitis Capital - General/AC - Digital/ACDAM Portfolio/Portfolio - Subscriptions"
)
SUBSCRIPTIONS_BASE_FALLBACK = Path(
    "/Users/gage/Amitis Group Holding Dropbox/Amitis Capital - General/AC - Digital/Portfolio - Subscriptions"
)
MANAGER_MATERIALS_BASE = Path(
    "/Users/gage/Amitis Group Holding Dropbox/Amitis Capital - General/AC - Digital/Manager Materials"
)
DROPBOX_SUBSCRIPTIONS_ROOT_GRANDLINE = (
    "/Amitis Capital - General/AC - Digital/ACDAM Portfolio/Portfolio - Subscriptions/GrandLine"
)
DROPBOX_SUBSCRIPTIONS_ROOT_GRANDLINE_FALLBACK = (
    "/Amitis Capital - General/AC - Digital/Portfolio - Subscriptions/GrandLine"
)
DROPBOX_MANAGER_MATERIALS_ROOT_GRANDLINE = (
    "/Amitis Capital - General/AC - Digital/Manager Materials/Grandline"
)
DROPBOX_IC_MATERIALS_ROOT_GRANDLINE = (
    "/Amitis Capital - General/AC - Digital/IC Materials/Grandline"
)
DROPBOX_SUBSCRIPTIONS_BASE = (
    "/Amitis Capital - General/AC - Digital/ACDAM Portfolio/Portfolio - Subscriptions"
)
DROPBOX_SUBSCRIPTIONS_BASE_FALLBACK = (
    "/Amitis Capital - General/AC - Digital/Portfolio - Subscriptions"
)
DROPBOX_MANAGER_MATERIALS_BASE = "/Amitis Capital - General/AC - Digital/Manager Materials"
DROPBOX_IC_MATERIALS_BASE = "/Amitis Capital - General/AC - Digital/IC Materials"

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".csv"}
LEGAL_DOC_TYPES = {"side_letter", "sub_agreement"}
DEFAULT_EXCLUDED_DOC_TYPES = "side_letter,sub_agreement"
EMBEDDING_MODEL = "text-embedding-3-small"
EMBEDDING_PRICE_PER_MILLION_TOKENS = 0.02

EXCLUSION_REGEX = re.compile(
    r"^(?:.*\bw-?9\b.*"
    r"|.*\bkyc\b.*"
    r"|.*\btax\b.*"
    r"|.*\bwire(?:[\s_-]?(?:instructions?|info))?\b.*"
    r"|.*\bform[\s_-]?adv\b.*"
    r"|~\$.*"
    r"|\.ds_store"
    r"|thumbs\.db)$",
    re.IGNORECASE,
)

DOC_TYPE_PATTERNS: list[tuple[str, str]] = [
    (r"(?:^|[^a-z0-9])ppm(?:$|[^a-z0-9])|private\s*placement\s*memorand|offering\s*memorand", "ppm"),
    (r"(?:^|[^a-z0-9])ddq(?:$|[^a-z0-9])|due\s*diligence\s*question", "ddq"),
    (r"factsheet|fact\s*sheet|\btearsheet\b|tear\s*sheet|one\s*pager|onepager", "factsheet"),
    (r"gross\s*return|net\s*returns?|daily\s*gross|pnl\s*daily|risk\s*contr[io]bution", "factsheet"),
    (r"fund\s*exposure|exposure", "exposure_report"),
    (r"reduced\s*fees?|fee\s*(calc|model)|fees?\s*calc", "fee_model"),
    (r"certificate\s*of\s*(change|incorporation)|incorporation", "corporate_doc"),
    (r"fund\s*structure|structure\s*chart", "structure_chart"),
    (r"\bic\s*report\b|investment\s*committee", "ic_report"),
    (r"\bodd\s*report\b|operational\s*due\s*diligence", "odd_report"),
    (
        r"monthly\s*(letter|update|commentary)"
        r"|quarterly\s*(letter|update|commentary)"
        r"|^[a-z]+\s*\d{4}\s*(letter|update)$",
        "letter",
    ),
    (r"audited|audit\s*report|\bfinancials?\b.*\d{4}", "audited_fs"),
    (r"side\s*letter", "side_letter"),
    (r"(sub|subscription)\s*(agreement|docs?|booklet|document)|payment\s*agreement", "sub_agreement"),
    (r"deck|pitch|investor\s*deck|presentation", "deck"),
    (r"(?:^|[^a-z0-9])lpa(?:$|[^a-z0-9])|limited\s*partnership\s*agreement", "lpa"),
    (r"\boperating\s*agreement\b|\bllc\s*agreement\b", "operating_agreement"),
]

DECISION_INDEXED = "indexed"
DECISION_SKIPPED_EXCLUDED = "skipped_excluded"
DECISION_SKIPPED_EXCLUDED_DOCTYPE = "skipped_excluded_doctype"
DECISION_SKIPPED_UNCHANGED = "skipped_unchanged"
DECISION_SKIPPED_UNSUPPORTED = "skipped_unsupported_extension"
DECISION_SKIPPED_DUPLICATE = "skipped_duplicate_format"
DECISION_FAILED = "failed"
DECISION_SOFT_DELETED = "soft_deleted"
TERMINAL_RESUME_DECISIONS = {
    DECISION_INDEXED,
    DECISION_SKIPPED_EXCLUDED,
    DECISION_SKIPPED_EXCLUDED_DOCTYPE,
    DECISION_SKIPPED_UNCHANGED,
    DECISION_SKIPPED_UNSUPPORTED,
    DECISION_SKIPPED_DUPLICATE,
}


@dataclass
class SourceRoot:
    path: str
    source_root: str
    provider: str


@dataclass
class SourceFile:
    path: Path
    source_root: str
    byte_size: int
    extension: str
    doc_type: str | None = None
    content_hash: str | None = None
    is_authoritative: bool = True
    provider: str = "local"
    dropbox_id: str | None = None
    dropbox_rev: str | None = None
    dropbox_content_hash: str | None = None

    @property
    def filepath(self) -> str:
        return str(self.path)

    @property
    def filename(self) -> str:
        return self.path.name


@dataclass
class ExtractedBlock:
    locator_kind: str
    locator_value: str | None
    text: str
    is_table_chunk: bool = False


@dataclass
class Chunk:
    chunk_index: int
    locator_kind: str
    locator_value: str | None
    text: str
    token_count: int
    is_table_chunk: bool
    embedding: list[float] | None = None


@dataclass
class PreparedDocument:
    source: SourceFile
    page_count: int | None
    chunks: list[Chunk]
    warnings: list[str]


class TokenCounter:
    def __init__(self) -> None:
        self._encoding = None
        try:
            import tiktoken

            self._encoding = tiktoken.get_encoding("cl100k_base")
        except Exception:
            print(
                "WARN: tiktoken unavailable; using approximate token counts. Install tiktoken before production runs.",
                file=sys.stderr,
            )

    @property
    def exact(self) -> bool:
        return self._encoding is not None

    def count(self, text: str) -> int:
        if not text:
            return 0
        if self._encoding:
            return len(self._encoding.encode(text))
        return max(1, len(text) // 4)

    def split(self, text: str, max_tokens: int, overlap_tokens: int) -> list[str]:
        text = text.strip()
        if not text:
            return []
        if self.count(text) <= max_tokens:
            return [text]

        if self._encoding:
            tokens = self._encoding.encode(text)
            chunks: list[str] = []
            step = max(1, max_tokens - overlap_tokens)
            start = 0
            while start < len(tokens):
                end = min(len(tokens), start + max_tokens)
                segment = self._encoding.decode(tokens[start:end]).strip()
                segment = trim_to_boundary(segment)
                if segment:
                    chunks.append(segment)
                if end >= len(tokens):
                    break
                start += step
            return chunks

        char_max = max_tokens * 4
        char_overlap = overlap_tokens * 4
        chunks = []
        start = 0
        while start < len(text):
            end = min(len(text), start + char_max)
            segment = trim_to_boundary(text[start:end]).strip()
            if segment:
                chunks.append(segment)
            if end >= len(text):
                break
            start += max(1, char_max - char_overlap)
        return chunks


class RunLog:
    def __init__(self, db_path: Path, fund_slug: str, argv: list[str], resume: bool) -> None:
        self.db_path = db_path
        self.fund_slug = fund_slug
        self.argv = argv
        self.conn = sqlite3.connect(str(db_path))
        self.conn.row_factory = sqlite3.Row
        init_sqlite(self.conn)
        self.run_id = self._find_resume_run() if resume else None
        self.resumed = self.run_id is not None
        if not self.run_id:
            self.run_id = uuid.uuid4().hex
            self.conn.execute(
                """
                INSERT INTO runs (run_id, fund_slug, started_at, cli_args)
                VALUES (?, ?, ?, ?)
                """,
                (self.run_id, fund_slug, utc_now(), json.dumps(argv)),
            )
            self.conn.commit()

    def _find_resume_run(self) -> str | None:
        row = self.conn.execute(
            """
            SELECT run_id FROM runs
            WHERE fund_slug = ? AND exit_status IS NULL
            ORDER BY started_at DESC LIMIT 1
            """,
            (self.fund_slug,),
        ).fetchone()
        return str(row["run_id"]) if row else None

    def existing_terminal_decisions(self) -> dict[str, str]:
        rows = self.conn.execute(
            """
            SELECT filepath, decision FROM file_events
            WHERE run_id = ? AND decision IN ({})
            """.format(",".join("?" for _ in TERMINAL_RESUME_DECISIONS)),
            (self.run_id, *sorted(TERMINAL_RESUME_DECISIONS)),
        ).fetchall()
        return {str(row["filepath"]): str(row["decision"]) for row in rows}

    def event(
        self,
        source: SourceFile | None,
        *,
        filepath: str | None = None,
        filename: str | None = None,
        source_root: str | None = None,
        byte_size: int | None = None,
        content_hash: str | None = None,
        doc_type: str | None = None,
        decision: str,
        chunks_produced: int | None = None,
        extraction_warnings: str | None = None,
        error_message: str | None = None,
        duration_ms: int | None = None,
    ) -> None:
        if source is not None:
            filepath = source.filepath
            filename = source.filename
            source_root = source.source_root
            byte_size = source.byte_size
            content_hash = content_hash if content_hash is not None else source.content_hash
            doc_type = doc_type if doc_type is not None else source.doc_type

        self.conn.execute(
            """
            INSERT INTO file_events (
              run_id, fund_slug, filepath, filename, source_root, byte_size,
              content_hash, doc_type, decision, chunks_produced,
              extraction_warnings, error_message, duration_ms, recorded_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                self.run_id,
                self.fund_slug,
                filepath or "",
                filename or "",
                source_root or "",
                byte_size,
                content_hash,
                doc_type,
                decision,
                chunks_produced,
                extraction_warnings,
                error_message,
                duration_ms,
                utc_now(),
            ),
        )
        self.conn.commit()

    def finish(
        self,
        *,
        exit_status: int,
        files_seen: int,
        preflight_tokens: int | None = None,
        preflight_cost: float | None = None,
        embedding_tokens: int = 0,
        embedding_cost: float = 0.0,
        notes: str | None = None,
    ) -> None:
        counts = {
            str(row["decision"]): int(row["count"])
            for row in self.conn.execute(
                "SELECT decision, COUNT(*) AS count FROM file_events WHERE run_id = ? GROUP BY decision",
                (self.run_id,),
            ).fetchall()
        }
        chunks_emitted = self.conn.execute(
            "SELECT COALESCE(SUM(chunks_produced), 0) AS chunks FROM file_events WHERE run_id = ? AND decision = ?",
            (self.run_id, DECISION_INDEXED),
        ).fetchone()["chunks"]
        files_skipped = sum(count for decision, count in counts.items() if decision.startswith("skipped_"))
        self.conn.execute(
            """
            UPDATE runs SET
              finished_at = ?,
              exit_status = ?,
              preflight_estimated_tokens = COALESCE(?, preflight_estimated_tokens),
              preflight_estimated_cost_usd = COALESCE(?, preflight_estimated_cost_usd),
              files_seen = ?,
              files_indexed = ?,
              files_skipped = ?,
              files_failed = ?,
              files_soft_deleted = ?,
              chunks_emitted = ?,
              embedding_tokens = ?,
              embedding_cost_usd = ?,
              notes = ?
            WHERE run_id = ?
            """,
            (
                utc_now(),
                exit_status,
                preflight_tokens,
                preflight_cost,
                files_seen,
                counts.get(DECISION_INDEXED, 0),
                files_skipped,
                counts.get(DECISION_FAILED, 0),
                counts.get(DECISION_SOFT_DELETED, 0),
                chunks_emitted,
                embedding_tokens,
                embedding_cost,
                notes,
                self.run_id,
            ),
        )
        self.conn.commit()

    def close(self) -> None:
        self.conn.close()


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def init_sqlite(conn: sqlite3.Connection) -> None:
    conn.executescript(
        """
        CREATE TABLE IF NOT EXISTS runs (
          run_id              TEXT PRIMARY KEY,
          fund_slug           TEXT NOT NULL,
          started_at          TEXT NOT NULL,
          finished_at         TEXT,
          exit_status         INTEGER,
          cli_args            TEXT NOT NULL,
          preflight_estimated_tokens INTEGER,
          preflight_estimated_cost_usd REAL,
          files_seen          INTEGER NOT NULL DEFAULT 0,
          files_indexed       INTEGER NOT NULL DEFAULT 0,
          files_skipped       INTEGER NOT NULL DEFAULT 0,
          files_failed        INTEGER NOT NULL DEFAULT 0,
          files_soft_deleted  INTEGER NOT NULL DEFAULT 0,
          chunks_emitted      INTEGER NOT NULL DEFAULT 0,
          embedding_tokens    INTEGER NOT NULL DEFAULT 0,
          embedding_cost_usd  REAL NOT NULL DEFAULT 0.0,
          notes               TEXT
        );

        CREATE INDEX IF NOT EXISTS idx_runs_fund_started ON runs(fund_slug, started_at DESC);
        CREATE INDEX IF NOT EXISTS idx_runs_unfinished ON runs(fund_slug, exit_status) WHERE exit_status IS NULL;

        CREATE TABLE IF NOT EXISTS file_events (
          id                  INTEGER PRIMARY KEY AUTOINCREMENT,
          run_id              TEXT NOT NULL REFERENCES runs(run_id),
          fund_slug           TEXT NOT NULL,
          filepath            TEXT NOT NULL,
          filename            TEXT NOT NULL,
          source_root         TEXT NOT NULL,
          byte_size           INTEGER,
          content_hash        TEXT,
          doc_type            TEXT,
          decision            TEXT NOT NULL,
          chunks_produced     INTEGER,
          extraction_warnings TEXT,
          error_message       TEXT,
          duration_ms         INTEGER,
          recorded_at         TEXT NOT NULL
        );

        CREATE INDEX IF NOT EXISTS idx_file_events_run ON file_events(run_id);
        CREATE INDEX IF NOT EXISTS idx_file_events_decision ON file_events(fund_slug, decision);
        CREATE INDEX IF NOT EXISTS idx_file_events_filepath ON file_events(fund_slug, filepath, recorded_at DESC);
        """
    )
    conn.commit()


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Index one fund's documents into Fund Doc Search.")
    parser.add_argument("--fund", required=True, help="Fund slug, e.g. grandline.")
    parser.add_argument("--display-name", help="Display name for new fund_managers rows.")
    parser.add_argument("--source-root", action="append", help="Fixture/custom root; may be repeated.")
    parser.add_argument("--source-root-subs", help="Subscriptions root override.")
    parser.add_argument("--source-root-mm", help="Manager Materials root override.")
    parser.add_argument("--source-root-ic", help="IC Materials root override.")
    parser.add_argument(
        "--source-provider",
        choices=("auto", "local", "dropbox"),
        default=os.environ.get("FUND_DOC_SOURCE_PROVIDER", "auto"),
        help="Read from local filesystem or Dropbox API. auto uses Dropbox when Dropbox credentials are set.",
    )
    parser.add_argument("--max-cost-usd", type=float, default=5.00)
    parser.add_argument("--exclude-doc-types", default=DEFAULT_EXCLUDED_DOC_TYPES)
    parser.add_argument("--keep-deleted", action="store_true")
    parser.add_argument("--skip-legal-name-prompt", action="store_true")
    parser.add_argument("--force-reembed", action="append", default=[])
    parser.add_argument("--resume", action="store_true")
    parser.add_argument("--embedding-model", default=EMBEDDING_MODEL)
    parser.add_argument("--anthropic-model", default=os.environ.get("ANTHROPIC_MODEL", "claude-sonnet-4-5-20250929"))
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    if load_dotenv:
        load_dotenv(PROJECT_ROOT / ".env.local")

    DATA_DIR.mkdir(exist_ok=True)
    LOCK_DIR.mkdir(exist_ok=True)
    (DATA_DIR / "test-fixtures" / "grandline-smoke").mkdir(parents=True, exist_ok=True)

    excluded_doc_types = parse_csv_arg(args.exclude_doc_types)
    force_reembed = {normalize_path_arg(p) for p in args.force_reembed}
    files_seen = 0
    preflight_tokens: int | None = None
    preflight_cost: float | None = None
    embedding_tokens = 0

    run_log = RunLog(LOG_DB_PATH, args.fund, sys.argv[1:], args.resume)
    if run_log.resumed:
        print(f"Resuming unfinished run_id={run_log.run_id}", file=sys.stderr)

    try:
        with acquire_lock(args.fund):
            roots = build_source_roots(args)
            raw_files = discover_files(roots)
            files_seen = len(raw_files)
            supabase = make_supabase_client()
            fund = get_or_create_fund_manager(supabase, args.fund, args.display_name)
            existing_docs = load_existing_documents(supabase, fund["id"])
            resume_terminal = run_log.existing_terminal_decisions() if args.resume else {}

            candidates, active_filepaths = filter_and_classify_files(
                raw_files=raw_files,
                excluded_doc_types=excluded_doc_types,
                run_log=run_log,
                resume_terminal=resume_terminal,
            )
            maybe_extract_legal_name(args, supabase, fund, candidates, excluded_doc_types)

            token_counter = TokenCounter()
            prepared: list[PreparedDocument] = []
            for source in progress(candidates, "Extracting"):
                if source.filepath in resume_terminal:
                    continue
                started = time.time()
                try:
                    with materialize_source_file(source) as local_path:
                        source.content_hash = sha256_file(local_path)
                        if source.byte_size == 0:
                            raise RuntimeError("empty_file")

                        existing = existing_docs.get(source.filepath)
                        unchanged = (
                            existing
                            and existing.get("content_hash") == source.content_hash
                            and not existing.get("deleted_at")
                            and source.filepath not in force_reembed
                        )
                        if unchanged:
                            run_log.event(
                                source,
                                decision=DECISION_SKIPPED_UNCHANGED,
                                duration_ms=elapsed_ms(started),
                            )
                            continue

                        blocks, page_count, warnings = extract_file(local_path)
                        chunks = chunk_blocks(blocks, token_counter)
                    if not chunks:
                        raise RuntimeError("no_extractable_text")
                    prepared.append(
                        PreparedDocument(
                            source=source,
                            page_count=page_count,
                            chunks=chunks,
                            warnings=warnings,
                        )
                    )
                except Exception as exc:
                    run_log.event(
                        source,
                        decision=DECISION_FAILED,
                        error_message=compact_error(exc),
                        duration_ms=elapsed_ms(started),
                    )

            preflight_tokens = sum(chunk.token_count for doc in prepared for chunk in doc.chunks)
            preflight_cost = preflight_tokens / 1_000_000 * EMBEDDING_PRICE_PER_MILLION_TOKENS
            print(
                f"Estimated embedding cost: ${preflight_cost:.4f} for {preflight_tokens:,} tokens "
                f"(ceiling ${args.max_cost_usd:.2f})",
                file=sys.stderr,
            )
            if preflight_cost > args.max_cost_usd:
                print(
                    f"Estimated cost ${preflight_cost:.4f} exceeds ceiling ${args.max_cost_usd:.2f}. "
                    "Use --max-cost-usd to override or narrow scope.",
                    file=sys.stderr,
                )
                run_log.finish(
                    exit_status=1,
                    files_seen=files_seen,
                    preflight_tokens=preflight_tokens,
                    preflight_cost=preflight_cost,
                    notes="preflight_cost_exceeded",
                )
                return 1

            for doc in progress(prepared, "Embedding/upserting"):
                started = time.time()
                try:
                    vectors, used_tokens = embed_chunks(doc.chunks, args.embedding_model)
                    embedding_tokens += used_tokens
                    for chunk, vector in zip(doc.chunks, vectors, strict=True):
                        chunk.embedding = vector
                    replace_document_chunks(supabase, fund["id"], doc, args.embedding_model)
                    run_log.event(
                        doc.source,
                        decision=DECISION_INDEXED,
                        chunks_produced=len(doc.chunks),
                        extraction_warnings="; ".join(doc.warnings) if doc.warnings else None,
                        duration_ms=elapsed_ms(started),
                    )
                except Exception as exc:
                    run_log.event(
                        doc.source,
                        decision=DECISION_FAILED,
                        error_message=compact_error(exc),
                        duration_ms=elapsed_ms(started),
                    )

            if not args.keep_deleted:
                soft_deleted = soft_delete_missing(supabase, fund["id"], active_filepaths)
                existing_active = [
                    row
                    for row in existing_docs.values()
                    if not row.get("deleted_at") and row.get("filepath") not in active_filepaths
                ]
                for row in existing_active:
                    run_log.event(
                        None,
                        filepath=row.get("filepath"),
                        filename=Path(row.get("filepath", "")).name,
                        source_root=row.get("source_root") or "",
                        byte_size=row.get("byte_size"),
                        content_hash=row.get("content_hash"),
                        doc_type=row.get("doc_type"),
                        decision=DECISION_SOFT_DELETED,
                    )
                if soft_deleted and soft_deleted != len(existing_active):
                    print(
                        f"WARN: soft-delete RPC reported {soft_deleted}, local log expected {len(existing_active)}.",
                        file=sys.stderr,
                    )

            embedding_cost = embedding_tokens / 1_000_000 * EMBEDDING_PRICE_PER_MILLION_TOKENS
            run_log.finish(
                exit_status=0,
                files_seen=files_seen,
                preflight_tokens=preflight_tokens,
                preflight_cost=preflight_cost,
                embedding_tokens=embedding_tokens,
                embedding_cost=embedding_cost,
            )
            print_run_summary(run_log)
            return 0
    except BlockingIOError:
        print(
            f"Lock held for fund={args.fund}; another indexer is running. "
            f"Lockfile: {LOCK_DIR / f'fund_{args.fund}.lock'}",
            file=sys.stderr,
        )
        run_log.finish(exit_status=1, files_seen=files_seen, notes="lock_held")
        return 1
    except Exception as exc:
        print(f"ERROR: {compact_error(exc)}", file=sys.stderr)
        run_log.finish(exit_status=1, files_seen=files_seen, notes=compact_error(exc))
        return 1
    finally:
        run_log.close()


def acquire_lock(fund_slug: str):
    lock_path = LOCK_DIR / f"fund_{fund_slug}.lock"
    lock_path.parent.mkdir(parents=True, exist_ok=True)
    handle = open(lock_path, "a+", encoding="utf-8")
    try:
        fcntl.flock(handle.fileno(), fcntl.LOCK_EX | fcntl.LOCK_NB)
    except BlockingIOError:
        handle.close()
        raise

    class LockContext:
        def __enter__(self):
            return handle

        def __exit__(self, exc_type, exc, tb):
            fcntl.flock(handle.fileno(), fcntl.LOCK_UN)
            handle.close()

    return LockContext()


def make_supabase_client():
    try:
        from supabase import create_client
    except Exception as exc:
        raise RuntimeError("Missing dependency: supabase. Install with pip install supabase.") from exc

    url = os.environ.get("NEXT_PUBLIC_SUPABASE_URL") or os.environ.get("SUPABASE_URL")
    key = os.environ.get("SUPABASE_SERVICE_ROLE_KEY")
    if not url or not key:
        raise RuntimeError("Missing NEXT_PUBLIC_SUPABASE_URL/SUPABASE_URL or SUPABASE_SERVICE_ROLE_KEY.")
    return create_client(url, key)


def get_or_create_fund_manager(supabase, slug: str, display_name: str | None) -> dict[str, Any]:
    resp = supabase.table("fund_managers").select("*").eq("slug", slug).limit(1).execute()
    if resp.data:
        return dict(resp.data[0])
    name = display_name or slug.replace("-", " ").replace("_", " ").title()
    resp = (
        supabase.table("fund_managers")
        .insert({"slug": slug, "display_name": name})
        .execute()
    )
    if not resp.data:
        raise RuntimeError(f"Unable to create fund_managers row for slug={slug}")
    return dict(resp.data[0])


def load_existing_documents(supabase, fund_id: str) -> dict[str, dict[str, Any]]:
    resp = (
        supabase.table("fund_documents")
        .select("id, filepath, filename, source_root, doc_type, content_hash, byte_size, deleted_at")
        .eq("fund_id", fund_id)
        .execute()
    )
    return {str(row["filepath"]): dict(row) for row in (resp.data or [])}


def replace_document_chunks(supabase, fund_id: str, doc: PreparedDocument, embedding_model: str) -> None:
    payload = []
    for chunk in doc.chunks:
        if chunk.embedding is None:
            raise RuntimeError("embedding_missing")
        payload.append(
            {
                "chunk_index": chunk.chunk_index,
                "locator_kind": chunk.locator_kind,
                "locator_value": chunk.locator_value,
                "is_table_chunk": chunk.is_table_chunk,
                "token_count": chunk.token_count,
                "text": chunk.text,
                "embedding": chunk.embedding,
                "embedding_model": embedding_model,
            }
        )
    supabase.rpc(
        "replace_fund_document_chunks",
        {
            "p_fund_id": fund_id,
            "p_filepath": doc.source.filepath,
            "p_filename": doc.source.filename,
            "p_source_root": doc.source.source_root,
            "p_doc_type": doc.source.doc_type,
            "p_content_hash": doc.source.content_hash,
            "p_is_authoritative": doc.source.is_authoritative,
            "p_byte_size": doc.source.byte_size,
            "p_page_count": doc.page_count,
            "p_chunks": payload,
        },
    ).execute()


def soft_delete_missing(supabase, fund_id: str, active_filepaths: set[str]) -> int:
    resp = supabase.rpc(
        "soft_delete_missing_fund_documents",
        {"p_fund_id": fund_id, "p_seen_filepaths": sorted(active_filepaths)},
    ).execute()
    if isinstance(resp.data, int):
        return resp.data
    return int(resp.data or 0)


def build_source_roots(args: argparse.Namespace) -> list[SourceRoot]:
    provider = resolve_source_provider(args)
    roots: list[SourceRoot] = []
    if args.source_root:
        for raw in args.source_root:
            roots.append(SourceRoot(raw, infer_source_root(raw), provider))
    if args.source_root_subs:
        roots.append(SourceRoot(args.source_root_subs, "subscriptions", provider))
    if args.source_root_mm:
        roots.append(SourceRoot(args.source_root_mm, "manager_materials", provider))
    if args.source_root_ic:
        roots.append(SourceRoot(args.source_root_ic, "ic_materials", provider))
    if roots:
        return [resolve_source_root(root) for root in roots]

    if args.fund.casefold() == "grandline":
        if provider == "dropbox":
            return [
                SourceRoot(
                    first_existing_dropbox_path(
                        DROPBOX_SUBSCRIPTIONS_ROOT_GRANDLINE,
                        DROPBOX_SUBSCRIPTIONS_ROOT_GRANDLINE_FALLBACK,
                    ),
                    "subscriptions",
                    provider,
                ),
                SourceRoot(DROPBOX_MANAGER_MATERIALS_ROOT_GRANDLINE, "manager_materials", provider),
                SourceRoot(DROPBOX_IC_MATERIALS_ROOT_GRANDLINE, "ic_materials", provider),
            ]
        return [
            SourceRoot(
                str(first_existing_path(SUBSCRIPTIONS_ROOT_GRANDLINE, SUBSCRIPTIONS_ROOT_GRANDLINE_FALLBACK)),
                "subscriptions",
                provider,
            ),
            SourceRoot(str(resolve_casefold_path(MANAGER_MATERIALS_ROOT_GRANDLINE)), "manager_materials", provider),
        ]

    display = args.display_name or args.fund.replace("-", " ").replace("_", " ").title().replace(" ", "")
    if provider == "dropbox":
        return [
            SourceRoot(
                first_existing_dropbox_path(
                    join_dropbox_path(DROPBOX_SUBSCRIPTIONS_BASE, display),
                    join_dropbox_path(DROPBOX_SUBSCRIPTIONS_BASE_FALLBACK, display),
                ),
                "subscriptions",
                provider,
            ),
            SourceRoot(join_dropbox_path(DROPBOX_MANAGER_MATERIALS_BASE, display), "manager_materials", provider),
            SourceRoot(join_dropbox_path(DROPBOX_IC_MATERIALS_BASE, display), "ic_materials", provider),
        ]
    return [
        SourceRoot(
            str(first_existing_path(SUBSCRIPTIONS_BASE / display, SUBSCRIPTIONS_BASE_FALLBACK / display)),
            "subscriptions",
            provider,
        ),
        SourceRoot(str(resolve_casefold_path(MANAGER_MATERIALS_BASE / display)), "manager_materials", provider),
    ]


def resolve_source_provider(args: argparse.Namespace) -> str:
    if args.source_provider != "auto":
        return args.source_provider
    dropbox_env = (
        os.environ.get("DROPBOX_MCP_TOKEN")
        or os.environ.get("DROPBOX_ACCESS_TOKEN")
        or (
            os.environ.get("DROPBOX_APP_KEY")
            and os.environ.get("DROPBOX_APP_SECRET")
            and os.environ.get("DROPBOX_REFRESH_TOKEN")
        )
    )
    return "dropbox" if dropbox_env else "local"


def resolve_source_root(root: SourceRoot) -> SourceRoot:
    if root.provider == "dropbox":
        return SourceRoot(normalize_dropbox_path(root.path), root.source_root, root.provider)
    return SourceRoot(str(resolve_casefold_path(Path(root.path).expanduser())), root.source_root, root.provider)


def infer_source_root(path: str | Path) -> str:
    lowered = str(path).casefold()
    if "ic material" in lowered or "investment committee" in lowered:
        return "ic_materials"
    if "portfolio - subscriptions" in lowered or "subscription" in lowered:
        return "subscriptions"
    return "manager_materials"


def resolve_casefold_path(path: Path) -> Path:
    path = path.expanduser()
    if path.exists():
        return path
    parts = path.parts
    if not parts:
        return path
    current = Path(parts[0])
    for part in parts[1:]:
        candidate = current / part
        if candidate.exists():
            current = candidate
            continue
        if current.exists() and current.is_dir():
            match = next((child for child in current.iterdir() if child.name.casefold() == part.casefold()), None)
            if match:
                current = match
                continue
        current = candidate
    return current


def first_existing_path(*paths: Path) -> Path:
    resolved = [resolve_casefold_path(path) for path in paths]
    return next((path for path in resolved if path.exists()), resolved[0])


def discover_files(roots: list[SourceRoot]) -> list[SourceFile]:
    files: list[SourceFile] = []
    for root in roots:
        if root.provider == "dropbox":
            files.extend(discover_dropbox_files(root))
            continue
        local_root = Path(root.path)
        label = root.source_root
        if not local_root.exists():
            print(f"WARN: source root does not exist: {local_root}", file=sys.stderr)
            continue
        for path in sorted((p for p in local_root.rglob("*") if p.is_file()), key=lambda p: str(p).casefold()):
            try:
                stat = path.stat()
                byte_size = stat.st_size
            except OSError:
                byte_size = None
            files.append(
                SourceFile(
                    path=path,
                    source_root=label,
                    byte_size=int(byte_size or 0),
                    extension=path.suffix.casefold(),
                    provider=root.provider,
                )
            )
    return files


def discover_dropbox_files(root: SourceRoot) -> list[SourceFile]:
    client = DropboxClient.from_env()
    entries = client.list_folder(root.path)
    files: list[SourceFile] = []
    for entry in entries:
        if entry.get(".tag") != "file":
            continue
        path_display = str(entry.get("path_display") or entry.get("path_lower") or "")
        if not path_display:
            continue
        files.append(
            SourceFile(
                path=Path(path_display),
                source_root=root.source_root,
                byte_size=int(entry.get("size") or 0),
                extension=Path(path_display).suffix.casefold(),
                provider=root.provider,
                dropbox_id=entry.get("id"),
                dropbox_rev=entry.get("rev"),
                dropbox_content_hash=entry.get("content_hash"),
            )
        )
    return sorted(files, key=lambda item: item.filepath.casefold())


@contextlib.contextmanager
def materialize_source_file(source: SourceFile):
    if source.provider == "dropbox":
        suffix = source.extension or source.path.suffix
        with tempfile.NamedTemporaryFile(prefix="fund-doc-", suffix=suffix) as handle:
            DropboxClient.from_env().download(source.filepath, handle.name)
            yield Path(handle.name)
        return

    if source.byte_size == 0:
        raise RuntimeError(
            "smart_sync_offline: file has byte_size=0; use --source-provider dropbox or make the Dropbox file available offline"
        )
    yield source.path


class DropboxClient:
    API = "https://api.dropboxapi.com/2"
    CONTENT = "https://content.dropboxapi.com/2"

    def __init__(self, token: str) -> None:
        self.token = token

    @classmethod
    def from_env(cls) -> "DropboxClient":
        token = os.environ.get("DROPBOX_ACCESS_TOKEN") or os.environ.get("DROPBOX_MCP_TOKEN")
        if os.environ.get("DROPBOX_REFRESH_TOKEN"):
            token = cls._refresh_access_token()
        if not token:
            raise RuntimeError(
                "Missing Dropbox credentials. Set DROPBOX_MCP_TOKEN, DROPBOX_ACCESS_TOKEN, "
                "or DROPBOX_APP_KEY/DROPBOX_APP_SECRET/DROPBOX_REFRESH_TOKEN."
            )
        return cls(token)

    @classmethod
    def _refresh_access_token(cls) -> str:
        app_key = os.environ.get("DROPBOX_APP_KEY")
        app_secret = os.environ.get("DROPBOX_APP_SECRET")
        refresh_token = os.environ.get("DROPBOX_REFRESH_TOKEN")
        if not app_key or not app_secret or not refresh_token:
            raise RuntimeError("Missing Dropbox OAuth refresh credentials.")
        body = urlencode({"grant_type": "refresh_token", "refresh_token": refresh_token}).encode("utf-8")
        req = request.Request("https://api.dropboxapi.com/oauth2/token", data=body, method="POST")
        req.add_header("Authorization", "Basic " + basic_auth(app_key, app_secret))
        req.add_header("Content-Type", "application/x-www-form-urlencoded")
        try:
            with request.urlopen(req, timeout=120) as resp:
                data = json.loads(resp.read().decode("utf-8"))
        except error.HTTPError as exc:
            raise RuntimeError(f"dropbox_token_refresh_failed:{exc.code}:{read_error_body(exc)}") from exc
        except error.URLError as exc:
            raise RuntimeError(f"dropbox_token_refresh_failed:{exc.reason}") from exc
        token = data.get("access_token")
        if not token:
            raise RuntimeError("dropbox_token_refresh_failed:missing_access_token")
        os.environ["DROPBOX_ACCESS_TOKEN"] = str(token)
        return str(token)

    def auth_headers(self) -> dict[str, str]:
        headers = {"Authorization": f"Bearer {self.token}"}
        namespace_id = os.environ.get("DROPBOX_ROOT_NAMESPACE_ID")
        if namespace_id:
            headers["Dropbox-API-Path-Root"] = json.dumps({".tag": "namespace_id", "namespace_id": namespace_id})
        return headers

    def list_folder(self, path: str) -> list[dict[str, Any]]:
        normalized = normalize_dropbox_path(path)
        payload = {"path": normalized, "recursive": True, "include_deleted": False, "limit": 2000}
        try:
            data = self._post_json(f"{self.API}/files/list_folder", payload)
        except RuntimeError as exc:
            if "path/not_found" in str(exc):
                print(f"WARN: Dropbox source root does not exist: {normalized}", file=sys.stderr)
                return []
            raise

        entries = list(data.get("entries") or [])
        while data.get("has_more"):
            data = self._post_json(f"{self.API}/files/list_folder/continue", {"cursor": data["cursor"]})
            entries.extend(data.get("entries") or [])
        return entries

    def download(self, path: str, output_path: str) -> None:
        headers = {
            **self.auth_headers(),
            "Dropbox-API-Arg": json.dumps({"path": normalize_dropbox_path(path)}),
        }
        req = request.Request(f"{self.CONTENT}/files/download", headers=headers, method="POST")
        try:
            with request.urlopen(req, timeout=120) as resp, open(output_path, "wb") as out:
                while True:
                    block = resp.read(1024 * 1024)
                    if not block:
                        break
                    out.write(block)
        except error.HTTPError as exc:
            raise RuntimeError(f"dropbox_download_failed:{exc.code}:{read_error_body(exc)}") from exc
        except error.URLError as exc:
            raise RuntimeError(f"dropbox_download_failed:{exc.reason}") from exc

    def _post_json(self, url: str, payload: dict[str, Any] | None) -> dict[str, Any]:
        body = b"null" if payload is None else json.dumps(payload).encode("utf-8")
        headers = {
            **self.auth_headers(),
            "Content-Type": "application/json",
        }
        req = request.Request(url, data=body, headers=headers, method="POST")
        try:
            with request.urlopen(req, timeout=120) as resp:
                return json.loads(resp.read().decode("utf-8"))
        except error.HTTPError as exc:
            raise RuntimeError(f"dropbox_api_failed:{exc.code}:{read_error_body(exc)}") from exc
        except error.URLError as exc:
            raise RuntimeError(f"dropbox_api_failed:{exc.reason}") from exc


def read_error_body(exc: error.HTTPError) -> str:
    try:
        body = exc.read().decode("utf-8", errors="replace")
    except Exception:
        body = str(exc)
    return re.sub(r"\s+", " ", body)[:500]


def basic_auth(username: str, password: str) -> str:
    return base64.b64encode(f"{username}:{password}".encode("utf-8")).decode("ascii")


def normalize_dropbox_path(path: str) -> str:
    value = str(path).strip()
    if not value or value == "/":
        return ""
    return "/" + value.strip("/")


def join_dropbox_path(*parts: str) -> str:
    joined = "/".join(part.strip("/") for part in parts if part)
    return normalize_dropbox_path(joined)


def first_existing_dropbox_path(*paths: str) -> str:
    client = DropboxClient.from_env()
    for path in paths:
        normalized = normalize_dropbox_path(path)
        if client.list_folder(normalized):
            return normalized
    return normalize_dropbox_path(paths[0])


def filter_and_classify_files(
    *,
    raw_files: list[SourceFile],
    excluded_doc_types: list[str],
    run_log: RunLog,
    resume_terminal: dict[str, str],
) -> tuple[list[SourceFile], set[str]]:
    supported: list[SourceFile] = []
    for source in raw_files:
        if source.filepath in resume_terminal:
            continue
        if EXCLUSION_REGEX.match(source.path.stem):
            run_log.event(source, decision=DECISION_SKIPPED_EXCLUDED)
            continue
        if source.extension not in SUPPORTED_EXTENSIONS:
            fmt = source.extension.lstrip(".") or "none"
            run_log.event(
                source,
                decision=DECISION_SKIPPED_UNSUPPORTED,
                error_message=f"format={fmt}",
            )
            continue
        source.doc_type = classify_doc_type(source.path.stem)
        if source.doc_type in excluded_doc_types:
            run_log.event(
                source,
                decision=DECISION_SKIPPED_EXCLUDED_DOCTYPE,
                error_message=f"doc_type={source.doc_type}",
            )
            continue
        supported.append(source)

    candidates = resolve_duplicate_formats(supported, run_log)
    active_filepaths = {source.filepath for source in candidates}
    active_filepaths.update(
        filepath
        for filepath, decision in resume_terminal.items()
        if decision in {DECISION_INDEXED, DECISION_SKIPPED_UNCHANGED}
    )
    return candidates, active_filepaths


def resolve_duplicate_formats(files: list[SourceFile], run_log: RunLog) -> list[SourceFile]:
    by_key: dict[tuple[str, str], list[SourceFile]] = {}
    for source in files:
        if source.extension in {".pdf", ".docx"}:
            by_key.setdefault(("all_roots", source.path.with_suffix("").name.casefold()), []).append(source)

    losers: set[str] = set()
    for group in by_key.values():
        has_pdf = any(item.extension == ".pdf" for item in group)
        has_docx = any(item.extension == ".docx" for item in group)
        if not (has_pdf and has_docx):
            continue
        doc_type = next((item.doc_type for item in group if item.doc_type), "other")
        preferred_ext = ".pdf" if doc_type in LEGAL_DOC_TYPES else ".docx"
        reason = "preferred_format=pdf_authoritative" if preferred_ext == ".pdf" else "preferred_format=docx"
        for source in group:
            if source.extension != preferred_ext:
                losers.add(source.filepath)
                run_log.event(source, decision=DECISION_SKIPPED_DUPLICATE, error_message=reason)

    candidates = []
    for source in files:
        if source.filepath not in losers:
            source.is_authoritative = True
            candidates.append(source)
    return candidates


def classify_doc_type(stem: str) -> str:
    lowered = stem.casefold()
    for pattern, doc_type in DOC_TYPE_PATTERNS:
        if re.search(pattern, lowered, flags=re.IGNORECASE):
            return doc_type
    return "other"


def maybe_extract_legal_name(
    args: argparse.Namespace,
    supabase,
    fund: dict[str, Any],
    candidates: list[SourceFile],
    excluded_doc_types: list[str],
) -> None:
    if fund.get("legal_name"):
        return

    ppm_candidates = [
        source
        for source in candidates
        if source.doc_type == "ppm" and source.extension == ".pdf" and source.byte_size > 0
    ]
    source = max(ppm_candidates, key=lambda item: item.byte_size, default=None)
    if source is None and not any(doc_type in excluded_doc_types for doc_type in LEGAL_DOC_TYPES):
        legal_candidates = [
            item
            for item in candidates
            if item.doc_type == "sub_agreement" and item.extension == ".pdf" and item.byte_size > 0
        ]
        source = max(legal_candidates, key=lambda item: item.byte_size, default=None)
    if source is None:
        print(
            "WARN: no locally available PPM PDF found for legal-name extraction; leaving legal_name empty. "
            "Dropbox online-only files must be made available offline first.",
            file=sys.stderr,
        )
        return

    try:
        with materialize_source_file(source) as local_path:
            text = extract_pdf_first_pages(local_path, max_pages=3)
        proposed = extract_legal_name_with_claude(text, args.anthropic_model)
    except Exception as exc:
        print(f"WARN: legal-name extraction failed: {compact_error(exc)}", file=sys.stderr)
        proposed = {"legal_name": None, "confidence": "low"}

    legal_name = proposed.get("legal_name")
    confidence = proposed.get("confidence") if proposed.get("confidence") in {"high", "medium", "low"} else "low"
    confirmed_at = None
    if args.skip_legal_name_prompt:
        print(f"WARN: legal name auto-extracted with confidence={confidence}; review pending.", file=sys.stderr)
    else:
        print("\nLegal-name extraction proposal:")
        print(json.dumps({"legal_name": legal_name, "confidence": confidence}, indent=2))
        edited = input("Confirm/edit legal_name (blank keeps proposal; 'skip' leaves empty): ").strip()
        if edited.casefold() == "skip":
            legal_name = None
        elif edited:
            legal_name = edited
        if legal_name:
            confirmed_at = utc_now()

    supabase.table("fund_managers").update(
        {
            "legal_name": legal_name,
            "legal_name_confidence": confidence,
            "legal_name_confirmed_at": confirmed_at,
        }
    ).eq("id", fund["id"]).execute()


def extract_legal_name_with_claude(text: str, model: str) -> dict[str, Any]:
    if not text.strip():
        raise RuntimeError("no text available for legal-name extraction")
    try:
        import anthropic
    except Exception as exc:
        raise RuntimeError("Missing dependency: anthropic. Install with pip install anthropic.") from exc
    if not os.environ.get("ANTHROPIC_API_KEY"):
        raise RuntimeError("Missing ANTHROPIC_API_KEY")
    client = anthropic.Anthropic()
    prompt = (
        "Extract the master fund legal name from the document excerpt. "
        "Return only JSON matching this schema: "
        '{"legal_name":"string or null","confidence":"high|medium|low"}.\n\n'
        f"Document excerpt:\n{text[:12000]}"
    )
    msg = client.messages.create(
        model=model,
        max_tokens=300,
        temperature=0,
        messages=[{"role": "user", "content": prompt}],
    )
    content = "".join(block.text for block in msg.content if getattr(block, "type", "") == "text")
    return parse_json_object(content)


def extract_pdf_first_pages(path: Path, max_pages: int) -> str:
    try:
        import pdfplumber
    except Exception as exc:
        raise RuntimeError("Missing dependency: pdfplumber. Install with pip install pdfplumber.") from exc
    texts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages[:max_pages]:
            texts.append(page.extract_text() or "")
    return "\n\n".join(texts).strip()


def extract_file(path: Path) -> tuple[list[ExtractedBlock], int | None, list[str]]:
    ext = path.suffix.casefold()
    if ext == ".pdf":
        return extract_pdf(path)
    if ext == ".docx":
        return extract_docx(path)
    if ext == ".xlsx":
        return extract_xlsx(path)
    if ext == ".csv":
        return extract_csv(path)
    raise RuntimeError(f"unsupported_extension:{ext}")


def extract_pdf(path: Path) -> tuple[list[ExtractedBlock], int | None, list[str]]:
    try:
        import pdfplumber
    except Exception as exc:
        raise RuntimeError("Missing dependency: pdfplumber. Install with pip install pdfplumber.") from exc
    blocks: list[ExtractedBlock] = []
    zero_text_pages = 0
    try:
        with pdfplumber.open(path) as pdf:
            for idx, page in enumerate(pdf.pages, start=1):
                text = (page.extract_text() or "").strip()
                if text:
                    blocks.append(ExtractedBlock("page", str(idx), text, False))
                else:
                    zero_text_pages += 1
                for table_idx, table in enumerate(page.extract_tables() or [], start=1):
                    table_text = table_to_text(table)
                    if table_text:
                        blocks.append(ExtractedBlock("page", str(idx), table_text, True))
            warnings = []
            if zero_text_pages:
                warnings.append(f"{zero_text_pages} pages had zero text after extraction")
            if not blocks:
                raise RuntimeError("no_extractable_text")
            return blocks, len(pdf.pages), warnings
    except Exception as exc:
        message = compact_error(exc)
        if "password" in message.casefold() or "encrypted" in message.casefold():
            raise RuntimeError("encrypted") from exc
        raise


def extract_docx(path: Path) -> tuple[list[ExtractedBlock], int | None, list[str]]:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("Missing dependency: python-docx. Install with pip install python-docx.") from exc
    doc = Document(path)
    blocks: list[ExtractedBlock] = []
    section = "document"
    buf: list[str] = []
    first_para_index = 0

    def flush() -> None:
        nonlocal buf, first_para_index
        text = "\n\n".join(part for part in buf if part.strip()).strip()
        if text:
            blocks.append(ExtractedBlock("section", section or f"para_{first_para_index}", text, False))
        buf = []

    for idx, para in enumerate(doc.paragraphs, start=1):
        text = para.text.strip()
        if not text:
            continue
        style_name = (para.style.name if para.style is not None else "").casefold()
        if style_name.startswith("heading"):
            flush()
            section = text[:180]
            first_para_index = idx
        else:
            if not buf:
                first_para_index = idx
            buf.append(text)
    flush()

    for table_idx, table in enumerate(doc.tables, start=1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        table_text = table_to_text(rows)
        if table_text:
            blocks.append(ExtractedBlock("table", f"{section}:table_{table_idx}", table_text, True))
    if not blocks:
        raise RuntimeError("no_extractable_text")
    return blocks, None, []


def extract_xlsx(path: Path) -> tuple[list[ExtractedBlock], int | None, list[str]]:
    try:
        import openpyxl
    except Exception as exc:
        raise RuntimeError("Missing dependency: openpyxl. Install with pip install openpyxl.") from exc
    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    blocks: list[ExtractedBlock] = []
    for sheet in workbook.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue
        non_empty_rows = [
            ["" if value is None else str(value) for value in row]
            for row in rows
            if any(value is not None and str(value).strip() for value in row)
        ]
        if not non_empty_rows:
            continue
        cells = sum(len(row) for row in non_empty_rows)
        group_size = 200 if cells > 5000 else len(non_empty_rows)
        for start in range(0, len(non_empty_rows), group_size):
            group = non_empty_rows[start : start + group_size]
            row_start = start + 1
            row_end = start + len(group)
            text = table_to_text(group)
            blocks.append(
                ExtractedBlock("sheet_row", f"{sheet.title}:row_{row_start}-row_{row_end}", text, True)
            )
    workbook.close()
    if not blocks:
        raise RuntimeError("no_extractable_text")
    return blocks, None, []


def extract_csv(path: Path) -> tuple[list[ExtractedBlock], int | None, list[str]]:
    blocks: list[ExtractedBlock] = []
    with open(path, "r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.reader(handle))
    if not rows:
        raise RuntimeError("no_extractable_text")
    header = rows[0]
    data_rows = rows[1:] if len(rows) > 1 else []
    if not data_rows:
        blocks.append(ExtractedBlock("row_range", "row_1-row_1", table_to_text([header]), True))
        return blocks, None, []
    group_size = 200
    for start in range(0, len(data_rows), group_size):
        group = [header] + data_rows[start : start + group_size]
        row_start = start + 1
        row_end = start + len(group) - 1
        blocks.append(ExtractedBlock("row_range", f"row_{row_start}-row_{row_end}", table_to_text(group), True))
    return blocks, None, []


def table_to_text(table: Iterable[Iterable[Any]]) -> str:
    rows: list[str] = []
    for row in table:
        cells = [clean_cell(cell) for cell in row]
        if any(cells):
            rows.append(" | ".join(cells))
    return "\n".join(rows).strip()


def clean_cell(value: Any) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def chunk_blocks(blocks: list[ExtractedBlock], token_counter: TokenCounter) -> list[Chunk]:
    chunks: list[Chunk] = []
    for block in blocks:
        text = normalize_text(block.text)
        if not text:
            continue
        if block.is_table_chunk:
            pieces = split_table_block(text, token_counter)
        else:
            pieces = token_counter.split(text, max_tokens=600, overlap_tokens=100)
        for piece in pieces:
            chunks.append(
                Chunk(
                    chunk_index=len(chunks),
                    locator_kind=block.locator_kind if block.locator_kind else "none",
                    locator_value=block.locator_value,
                    text=piece,
                    token_count=token_counter.count(piece),
                    is_table_chunk=block.is_table_chunk,
                )
            )
    return chunks


def split_table_block(text: str, token_counter: TokenCounter) -> list[str]:
    if token_counter.count(text) <= 4000:
        return [text]
    rows = text.splitlines()
    if len(rows) <= 1:
        return token_counter.split(text, max_tokens=4000, overlap_tokens=0)
    header = rows[0]
    chunks: list[str] = []
    current = [header]
    for row in rows[1:]:
        candidate = "\n".join(current + [row])
        if token_counter.count(candidate) > 4000 and len(current) > 1:
            chunks.append("\n".join(current))
            current = [header, row]
        else:
            current.append(row)
    if len(current) > 1:
        chunks.append("\n".join(current))
    return chunks


def normalize_text(text: str) -> str:
    return re.sub(r"\n{3,}", "\n\n", text.replace("\x00", "")).strip()


def trim_to_boundary(text: str) -> str:
    if len(text) < 500:
        return text
    min_pos = int(len(text) * 0.65)
    boundary_positions = [
        text.rfind("\n\n", min_pos),
        text.rfind(". ", min_pos),
        text.rfind("! ", min_pos),
        text.rfind("? ", min_pos),
    ]
    pos = max(boundary_positions)
    if pos > min_pos:
        return text[: pos + 1]
    return text


def embed_chunks(chunks: list[Chunk], model: str) -> tuple[list[list[float]], int]:
    try:
        from openai import OpenAI
    except Exception as exc:
        raise RuntimeError("Missing dependency: openai. Install with pip install openai.") from exc
    api_key = os.environ.get("OPENAI_API_KEY", "").strip()
    if not is_plausible_openai_key(api_key) and sys.stdin.isatty():
        api_key = getpass.getpass("Paste OpenAI API key and press Enter: ").strip()
        if api_key:
            os.environ["OPENAI_API_KEY"] = api_key
    if not is_plausible_openai_key(os.environ.get("OPENAI_API_KEY", "")):
        raise RuntimeError("Missing or invalid OPENAI_API_KEY")
    client = OpenAI()
    vectors: list[list[float]] = []
    used_tokens = 0
    texts = [chunk.text for chunk in chunks]
    for start in range(0, len(texts), 100):
        batch = texts[start : start + 100]
        resp = retry_api_call(lambda: client.embeddings.create(model=model, input=batch))
        vectors.extend(item.embedding for item in resp.data)
        usage = getattr(resp, "usage", None)
        used_tokens += int(getattr(usage, "prompt_tokens", 0) or 0)
    if used_tokens == 0:
        used_tokens = sum(chunk.token_count for chunk in chunks)
    return vectors, used_tokens


def retry_api_call(func):
    delay = 1.0
    last_exc: Exception | None = None
    for _ in range(5):
        try:
            return func()
        except Exception as exc:
            last_exc = exc
            time.sleep(delay)
            delay *= 2
    assert last_exc is not None
    raise last_exc


def is_plausible_openai_key(value: str) -> bool:
    key = value.strip()
    if not key.startswith(("sk-", "sk-proj-")):
        return False
    lowered = key.casefold()
    return "your-real-openai-key" not in lowered and "placeholder" not in lowered


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def parse_csv_arg(value: str | None) -> list[str]:
    if value is None:
        return []
    return [part.strip() for part in value.split(",") if part.strip()]


def normalize_path_arg(value: str) -> str:
    return str(Path(value).expanduser())


def parse_json_object(text: str) -> dict[str, Any]:
    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise RuntimeError(f"Could not parse JSON from model response: {text[:200]}")
    return json.loads(text[start : end + 1])


def compact_error(exc: BaseException) -> str:
    message = str(exc).strip() or exc.__class__.__name__
    return re.sub(r"\s+", " ", message)[:1000]


def elapsed_ms(started: float) -> int:
    return int((time.time() - started) * 1000)


def progress(items: list[Any], desc: str):
    if tqdm is not None:
        return tqdm(items, desc=desc, file=sys.stderr)
    return items


def print_run_summary(run_log: RunLog) -> None:
    row = run_log.conn.execute(
        """
        SELECT files_seen, files_indexed, files_skipped, files_failed,
               files_soft_deleted, chunks_emitted, embedding_tokens,
               embedding_cost_usd
        FROM runs WHERE run_id = ?
        """,
        (run_log.run_id,),
    ).fetchone()
    print(json.dumps({"run_id": run_log.run_id, **dict(row)}, indent=2))


if __name__ == "__main__":
    raise SystemExit(main())
